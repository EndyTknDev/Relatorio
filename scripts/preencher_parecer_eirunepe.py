from __future__ import annotations

import hashlib
from pathlib import Path
from zipfile import ZipFile

from docx import Document


ROOT = Path(r"C:\Users\endyt\projetos\relatorio-plugin")
TEMPLATE = Path(r"C:\Users\endyt\.codex\plugins\cache\relatorios-locais\relatorio-plugin\0.1.0\template\PARECER_TECNICO_DEFESA_CIVIL.docx")
OUTPUT = ROOT / "outputs" / "Fevereiro 2026" / "PARECER TECNICO DEFESA CIVIL - EIRUNEPE - 001-2026.docx"
EXPECTED_TEMPLATE_SHA256 = "C7E35046D4DF414E1CA78E1C0927C2C602BB59381D652990DF41F92315C9323F"


def sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest().upper()


def set_paragraph(paragraph, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for run in paragraph.runs[1:]:
        run.text = ""


def set_suffix(paragraph, prefix: str, value: str) -> None:
    full = "".join(run.text for run in paragraph.runs)
    if not full.startswith(prefix):
        raise ValueError(f"Prefixo inesperado: {full!r}; esperado {prefix!r}")
    remaining = len(prefix)
    placed = False
    for run in paragraph.runs:
        old = run.text
        if remaining >= len(old):
            remaining -= len(old)
            continue
        keep = old[:remaining]
        remaining = 0
        run.text = keep + value
        placed = True
        break
    if not placed:
        paragraph.runs[-1].text += value
    found_value = False
    consumed = 0
    for run in paragraph.runs:
        consumed += len(run.text)
        if found_value:
            run.text = ""
        elif value and value in run.text:
            found_value = True


def unique_cell(table, row: int, col: int):
    return table.rows[row].cells[col]


def package_image_inventory(path: Path) -> tuple[int, int, int]:
    with ZipFile(path) as zf:
        names = zf.namelist()
        media = sum(name.startswith("word/media/") and not name.endswith("/") for name in names)
        xml = "\n".join(
            zf.read(name).decode("utf-8", errors="ignore")
            for name in names
            if name.endswith((".xml", ".rels"))
        )
    return media, xml.count("relationships/image"), xml.count("<a:blip") + xml.count("<v:imagedata")


def structure(doc: Document) -> dict:
    return {
        "body_paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "rows": [len(table.rows) for table in doc.tables],
        "cells": [sum(len(row.cells) for row in table.rows) for table in doc.tables],
        "cell_paragraphs": [
            [len(cell.paragraphs) for row in table.rows for cell in row.cells]
            for table in doc.tables
        ],
        "sections": len(doc.sections),
        "inline_shapes": len(doc.inline_shapes),
    }


def main() -> None:
    if sha256(TEMPLATE) != EXPECTED_TEMPLATE_SHA256:
        raise RuntimeError("Hash do template instalado diverge da referência.")
    if not OUTPUT.is_file():
        raise FileNotFoundError(OUTPUT)
    baseline_images = package_image_inventory(TEMPLATE)

    baseline = Document(TEMPLATE)
    doc = Document(OUTPUT)
    baseline_structure = structure(baseline)

    set_paragraph(doc.paragraphs[1], "PARECER TÉCNICO Nº 001/2026")
    set_paragraph(doc.paragraphs[2], "Assunto: Parecer técnico sobre os danos e impactos decorrentes da inundação causada pela cheia do Rio Juruá, para fundamentação da situação de emergência")
    set_paragraph(doc.paragraphs[3], "Desastre: Inundação")

    t0 = doc.tables[0]
    set_suffix(unique_cell(t0, 1, 0).paragraphs[0], "Município: ", "Eirunepé")
    set_suffix(unique_cell(t0, 1, 1).paragraphs[0], "UF: ", "Amazonas - AM")
    set_suffix(unique_cell(t0, 2, 0).paragraphs[0], "Decreto: ", "009/2026/GABPRE/PME, de 09 de fevereiro de 2026")
    set_suffix(unique_cell(t0, 2, 1).paragraphs[0], "Publicação: ", "Diário Oficial Eletrônico dos Municípios do Estado do Amazonas, edição 4041, de 09 de fevereiro de 2026, código 348C295F")
    set_suffix(unique_cell(t0, 3, 0).paragraphs[0], "Código COBRADE: ", "1.2.1.0.0")
    set_suffix(unique_cell(t0, 3, 1).paragraphs[0], "Tipo: ", "Inundação")
    set_suffix(unique_cell(t0, 3, 2).paragraphs[0], "Data: ", "09/02/2026")
    set_suffix(unique_cell(t0, 3, 3).paragraphs[0], "Hora: ", "10h00")

    causa = unique_cell(t0, 4, 0).paragraphs
    set_paragraph(causa[2], "No período de 04 a 09 de fevereiro de 2026, o aumento das chuvas e a elevação do Rio Juruá produziram inundação gradual em Eirunepé. Em 04/02/2026, o rio alcançou 16,89 metros, acima da cota de transbordamento, exigindo medidas urgentes de proteção, assistência e restabelecimento de serviços.")
    set_paragraph(causa[3], "O cenário foi acompanhado pelo Centro de Monitoramento e Alerta, em cooperação com o Serviço Geológico do Brasil, o SIPAM e outras instituições técnicas. O Alerta CEMADEN nº 0161/2026 indicou possibilidade de inundação gradual do Rio Juruá e de igarapés, enquanto relatórios municipais de Defesa Civil, Saúde, Assistência Social, Agricultura e Extensão Rural confirmaram impactos multissetoriais.")
    set_paragraph(causa[4], "Os levantamentos municipais registraram 4.931 pessoas atingidas na zona urbana e 6.550 na zona rural, totalizando 11.481 pessoas. Foram informados danos em 300 residências, cinco unidades de ensino e duas unidades de saúde, perda de produção agrícola familiar em 150 hectares, comprometimento de vias e sistemas essenciais de abastecimento, deslocamento de famílias, aumento da vulnerabilidade social e agravamento de riscos sanitários.")
    set_paragraph(causa[5], "As áreas urbanas com maior incidência foram os bairros Nossa Senhora do Perpétuo Socorro, especialmente a localidade Morada do Sol, Santo Antônio, São José e Nossa Senhora Aparecida. Na zona rural, foram afetadas as comunidades e regiões do Alto Juruá, Baixo Juruá, Rio Gregório, Rio Eiru e Rio Tarauacá. A publicação oficial consultada não informa distâncias ou quilometragens verificáveis para essas localidades.")

    t1 = doc.tables[1]
    localidades = unique_cell(t1, 0, 0).paragraphs
    set_paragraph(localidades[0], "Áreas urbanas: bairros Nossa Senhora do Perpétuo Socorro, com destaque para Morada do Sol, Santo Antônio, São José e Nossa Senhora Aparecida.")
    set_paragraph(localidades[1], "Áreas rurais: comunidades e regiões do Alto Juruá, Baixo Juruá, Rio Gregório, Rio Eiru e Rio Tarauacá.")
    set_paragraph(localidades[2], "Distâncias e quilometragens: não informadas na publicação oficial consultada.")
    set_paragraph(unique_cell(t1, 1, 0).paragraphs[0], "Situação de anormalidade: Situação de emergência")
    set_paragraph(unique_cell(t1, 1, 1).paragraphs[0], "Desastre Nível: II")
    set_paragraph(unique_cell(t1, 2, 0).paragraphs[0], "Protocolo de Registro no S2ID: ainda não registrado")

    t2 = doc.tables[2]
    humanos = unique_cell(t2, 0, 0).paragraphs
    set_paragraph(humanos[2], "Os levantamentos técnicos contabilizaram 11.481 pessoas atingidas, sendo 4.931 na zona urbana e 6.550 na zona rural. Houve danos em 300 residências e deslocamento de famílias, com maior gravidade em Morada do Sol. A fonte oficial consultada não apresenta número consolidado de famílias, e a capacidade operacional, administrativa e orçamentária municipal foi considerada insuficiente para resposta isolada.")
    set_paragraph(humanos[3], "A população atingida inclui moradores de áreas ribeirinhas e agricultores familiares. A perda de produção em 150 hectares compromete alimentos, trabalho e renda, ampliando a vulnerabilidade social e as dificuldades de subsistência das famílias afetadas.")
    set_paragraph(humanos[4], "O alagamento de residências e o comprometimento das condições mínimas de subsistência elevaram os riscos sanitários. A população deslocada e os moradores expostos à água da inundação demandam atendimento, monitoramento sanitário e medidas preventivas para redução de riscos epidemiológicos.")

    materiais = unique_cell(t2, 1, 0).paragraphs
    set_paragraph(materiais[2], "A inundação atingiu residências, vias de acesso, sistemas essenciais de abastecimento, cinco unidades de ensino e duas unidades de saúde. A interrupção ou limitação desses serviços e o contato com água potencialmente contaminada ampliam os riscos à mobilidade, à continuidade do atendimento público e à saúde coletiva.")
    set_paragraph(materiais[3], "Os danos materiais e econômicos informados abrangem 300 residências, cinco unidades de ensino, duas unidades de saúde e 150 hectares de produção agrícola familiar, com valores individualizados na relação a seguir.")
    set_paragraph(materiais[4], "O diagnóstico indica a necessidade de manter o atendimento às áreas urbanas e rurais, proteger moradias e bens, assegurar acessos provisórios e prover alimentos, água potável, combustível e kits de higiene. Conforme os dados fornecidos, os prejuízos e as aquisições necessárias são:")

    t3 = doc.tables[3]
    continuation = unique_cell(t3, 0, 0).paragraphs
    set_paragraph(continuation[0], "As medidas devem priorizar a assistência humanitária, o acesso às áreas inundadas, a continuidade de serviços essenciais e a redução dos riscos sanitários. Os valores abaixo correspondem às estimativas e aquisições informadas para a resposta municipal:")
    items = [
        "-Prejuízos com danos em 300 residências atingidas pela inundação - R$ 1.500.000,00;",
        "-Prejuízos com perda de produção agrícola familiar em 150 hectares - R$ 750.000,00;",
        "-Prejuízos com danos em 5 unidades de ensino - R$ 250.000,00;",
        "-Prejuízos com danos em 2 unidades de saúde - R$ 180.000,00.",
        "-Aquisição de 1.000 cestas de alimentos R$ 180.000,00.",
        "-Aquisição de 20.000 litros de água mineral R$ 40.000,00.",
        "-Aquisição de 2.000 litros de gasolina R$ 14.000,00.",
        "-Aquisição de 1.000 kits de higiene pessoal R$ 60.000,00.",
    ]
    for paragraph, item in zip(continuation[1:9], items, strict=True):
        set_paragraph(paragraph, item)

    ambientais = unique_cell(t3, 1, 0).paragraphs
    set_paragraph(ambientais[1], "O monitoramento técnico registrou transbordamento do Rio Juruá e inundação de igarapés e áreas ribeirinhas, com abrangência urbana e rural. As ocorrências alcançaram bairros, comunidades, vias, moradias e equipamentos públicos, alterando as condições ambientais e de circulação nas áreas inundadas.")
    set_paragraph(ambientais[2], "A inundação causou perda de produção agrícola familiar em 150 hectares e expôs áreas produtivas e recursos hídricos a degradação e contaminação. Esse quadro se relaciona diretamente aos riscos sanitários decorrentes do contato com água de inundação e ao possível aumento de agravos de veiculação hídrica, exigindo monitoramento ambiental e epidemiológico contínuo.")

    t4 = doc.tables[4]
    humanos_inst = unique_cell(t4, 0, 0).paragraphs
    set_paragraph(humanos_inst[2], "Sob coordenação da Secretaria Municipal de Proteção e Defesa Civil, foram mobilizadas equipes municipais de Defesa Civil, Saúde, Assistência Social, Agricultura, Extensão Rural e Obras. As ações confirmadas compreendem vistorias técnicas, levantamento de danos, atendimento às famílias afetadas, monitoramento sanitário, medidas preventivas contra riscos epidemiológicos, elevação de assoalhos, proteção de bens e construção de pontes e trapiches provisórios e definitivos.")

    materiais_rec = unique_cell(t4, 1, 0).paragraphs
    set_paragraph(materiais_rec[2], "Foram empregados recursos municipais na atuação de equipes, assistência humanitária, proteção de moradias e bens e manutenção de acessos. Para continuidade da resposta, foram informadas as aquisições de 1.000 cestas de alimentos, 20.000 litros de água mineral, 2.000 litros de gasolina e 1.000 kits de higiene pessoal.")
    set_paragraph(materiais_rec[3], "Não foi localizado, nas fontes oficiais consultadas, valor consolidado já executado pelo Município até 09/02/2026. Por isso, os valores apresentados neste parecer permanecem discriminados por item, sem serem tratados como despesas já realizadas.")
    set_paragraph(materiais_rec[4], "A magnitude do evento supera a capacidade de resposta isolada do Município, tornando necessário apoio técnico, logístico e financeiro dos Governos Estadual e Federal para assistência humanitária, reabilitação do cenário, restabelecimento de serviços essenciais e recuperação das áreas atingidas.")

    set_paragraph(doc.paragraphs[26], "Diante do exposto, os registros hidrológicos, os levantamentos municipais, a abrangência territorial, as 11.481 pessoas atingidas, os danos habitacionais, educacionais, sanitários e agrícolas, os prejuízos informados e a insuficiência da capacidade municipal de resposta caracterizam desastre de nível II e fundamentam a Situação de Emergência declarada pelo Decreto nº 009/2026/GABPRE/PME, bem como a remessa da documentação aos órgãos competentes para homologação e reconhecimento, sem pressupor reconhecimento já concedido.")
    set_paragraph(doc.paragraphs[27], "Eirunepé, 09 de fevereiro de 2026.")

    doc.save(OUTPUT)
    checked = Document(OUTPUT)
    if structure(checked) != baseline_structure:
        raise RuntimeError("A estrutura do documento foi alterada.")
    if package_image_inventory(OUTPUT) != baseline_images:
        raise RuntimeError("O inventário de imagens da saída diverge do template.")
    full_text = "\n".join(p.text for p in checked.paragraphs)
    full_text += "\n" + "\n".join(cell.text for table in checked.tables for row in table.rows for cell in row.cells)
    required = ["Eirunepé", "1.2.1.0.0", "11.481", "348C295F", "Rio Juruá"]
    missing = [item for item in required if item not in full_text]
    if missing:
        raise RuntimeError(f"Dados obrigatórios ausentes: {missing}")
    print(OUTPUT)
    print(f"template_sha256={sha256(TEMPLATE)}")
    print(f"structure={structure(checked)}")
    print(f"acara_occurrences={full_text.count('Acará')}")


if __name__ == "__main__":
    main()
