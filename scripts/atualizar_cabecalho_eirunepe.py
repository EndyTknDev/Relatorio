from __future__ import annotations

import hashlib
import os
import tempfile
from io import BytesIO
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile, ZipInfo

from lxml import etree
from PIL import Image
from docx import Document


ROOT = Path(r"C:\Users\endyt\projetos\relatorio-plugin")
DOCX = ROOT / "outputs" / "Fevereiro 2026" / "PARECER TECNICO DEFESA CIVIL - EIRUNEPE - 001-2026.docx"
SOURCE_IMAGE = Path(r"C:\Users\endyt\AppData\Local\Temp\codex-clipboard-09a10edb-cc21-474f-a4e9-d1636437ec38.png")
TEMPLATE = Path(r"C:\Users\endyt\.codex\plugins\cache\relatorios-locais\relatorio-plugin\0.1.0\template\PARECER_TECNICO_DEFESA_CIVIL.docx")
EXPECTED_TEMPLATE_SHA256 = "C7E35046D4DF414E1CA78E1C0927C2C602BB59381D652990DF41F92315C9323F"

NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


def sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest().upper()


def structure(path: Path) -> dict:
    doc = Document(path)
    return {
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "rows": [len(table.rows) for table in doc.tables],
        "cells": [sum(len(row.cells) for row in table.rows) for table in doc.tables],
        "cell_paragraphs": [
            [len(cell.paragraphs) for row in table.rows for cell in row.cells]
            for table in doc.tables
        ],
        "sections": len(doc.sections),
    }


def render_crop(source: Image.Image, crop: tuple[int, int, int, int], size: tuple[int, int]) -> bytes:
    region = source.crop(crop).convert("RGB")
    region.thumbnail((size[0] - 16, size[1] - 16), Image.Resampling.LANCZOS)
    canvas = Image.new("RGB", size, "white")
    x = (size[0] - region.width) // 2
    y = (size[1] - region.height) // 2
    canvas.paste(region, (x, y))
    output = BytesIO()
    canvas.save(output, format="JPEG", quality=95, subsampling=0)
    return output.getvalue()


def update_header(xml_bytes: bytes) -> bytes:
    parser = etree.XMLParser(remove_blank_text=False)
    root = etree.fromstring(xml_bytes, parser)
    text_nodes = root.xpath(".//w:txbxContent//w:t", namespaces=NS)
    for node in text_nodes:
        if node.text:
            node.text = node.text.replace("Acará", "Eirunepé")
            node.text = node.text.replace("Políticas", "Proteção")
            node.text = node.text.replace("Públicas Gabinete da Secretaria", "e Defesa Civil")
            if node.text == "da":
                node.text = "de"

    paragraphs = root.xpath(".//w:txbxContent/w:p", namespaces=NS)
    rendered = ["".join(node.text or "" for node in p.xpath(".//w:t", namespaces=NS)) for p in paragraphs]
    expected = [
        "Município de Eirunepé Prefeitura Municipal de Eirunepé",
        "Secretaria Municipal de Proteção e Defesa Civil",
        "Coordenadoria Municipal de Proteção e Defesa Civil",
    ] * 2
    if rendered != expected:
        raise RuntimeError(f"Texto inesperado no cabeçalho: {rendered!r}")
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)


def clone_info(info: ZipInfo) -> ZipInfo:
    cloned = ZipInfo(info.filename, date_time=info.date_time)
    cloned.compress_type = info.compress_type
    cloned.comment = info.comment
    cloned.extra = info.extra
    cloned.internal_attr = info.internal_attr
    cloned.external_attr = info.external_attr
    cloned.create_system = info.create_system
    return cloned


def main() -> None:
    if sha256(TEMPLATE) != EXPECTED_TEMPLATE_SHA256:
        raise RuntimeError("O template protegido foi alterado.")
    if not DOCX.is_file() or not SOURCE_IMAGE.is_file():
        raise FileNotFoundError("Documento ou imagem de origem não encontrado.")

    before_structure = structure(DOCX)
    before_hash = sha256(DOCX)
    source = Image.open(SOURCE_IMAGE)
    replacements = {
        "word/media/image1.jpeg": render_crop(source, (125, 88, 430, 369), (444, 300)),
        "word/media/image2.jpeg": render_crop(source, (100, 85, 450, 480), (371, 300)),
        "word/media/image3.jpeg": render_crop(source, (110, 365, 442, 472), (736, 300)),
    }

    with ZipFile(DOCX, "r") as source_zip:
        names = source_zip.namelist()
        if "word/header1.xml" not in names or any(name not in names for name in replacements):
            raise RuntimeError("Estrutura de cabeçalho inesperada.")
        header = update_header(source_zip.read("word/header1.xml"))
        fd, temp_name = tempfile.mkstemp(prefix="cabecalho_eirunepe_", suffix=".docx", dir=DOCX.parent)
        os.close(fd)
        temp_path = Path(temp_name)
        try:
            with ZipFile(temp_path, "w", compression=ZIP_DEFLATED) as target_zip:
                for info in source_zip.infolist():
                    data = source_zip.read(info.filename)
                    if info.filename == "word/header1.xml":
                        data = header
                    elif info.filename in replacements:
                        data = replacements[info.filename]
                    target_zip.writestr(clone_info(info), data)
            os.replace(temp_path, DOCX)
        finally:
            if temp_path.exists():
                temp_path.unlink()

    after_structure = structure(DOCX)
    if after_structure != before_structure:
        raise RuntimeError("A estrutura do documento foi alterada.")

    with ZipFile(DOCX) as zf:
        header = etree.fromstring(zf.read("word/header1.xml"))
        header_text = " ".join(header.xpath(".//w:txbxContent//w:t/text()", namespaces=NS))
        image_relationships = zf.read("word/_rels/header1.xml.rels").count(b"relationships/image")
        media_names = sorted(name for name in zf.namelist() if name.startswith("word/media/") and not name.endswith("/"))
        image_hashes = {name: hashlib.sha256(zf.read(name)).hexdigest() for name in media_names}
    if "Acará" in header_text or image_relationships != 3 or len(media_names) != 3:
        raise RuntimeError("A identidade antiga ou o inventário de imagens está incorreto.")
    if any(text not in header_text for text in ("Município de Eirunepé", "Prefeitura Municipal de Eirunepé", "Secretaria Municipal de Proteção e Defesa Civil", "Coordenadoria Municipal de Proteção e Defesa Civil")):
        raise RuntimeError("O cabeçalho não contém todas as unidades institucionais solicitadas.")
    print(DOCX)
    print(f"before_sha256={before_hash}")
    print(f"after_sha256={sha256(DOCX)}")
    print(f"template_sha256={sha256(TEMPLATE)}")
    print(f"structure={after_structure}")
    print(f"media={image_hashes}")


if __name__ == "__main__":
    main()
